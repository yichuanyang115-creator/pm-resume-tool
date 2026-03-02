"""
简历优化工具 - FastAPI 后端服务
调用 GLM-4-Flash 模型分析简历并生成优化版本
"""

import os
import json
import uuid
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# 文档解析
import pdfplumber
from docx import Document as DocxDocument

# 文档生成 - PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, KeepTogether
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

# 文档生成 - Word
from docx import Document as WordDocument
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# AI 模型
from zhipuai import ZhipuAI

# ────────────────────────────────────────
# 配置
# ────────────────────────────────────────
API_KEY = "df38ea4eb1584fabb0851837b55d1910.u4nwnhBu7526QS4e"
MODEL   = "glm-4.7-flash"

BASE_DIR   = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
STATIC_DIR = BASE_DIR / "static"

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# 注册中文 CID 字体（reportlab 内置，无需额外字体文件）
try:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    CH_FONT = "STSong-Light"
except Exception:
    CH_FONT = "Helvetica"  # Render 环境回退字体

# ────────────────────────────────────────
# FastAPI 应用
# ────────────────────────────────────────
app = FastAPI(title="简历优化工具")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 注意：静态文件由 Netlify 托管，后端仅提供 API


@app.get("/")
async def root():
    return {"status": "ok", "service": "简历优化工具 API", "version": "1.0.0"}


# ────────────────────────────────────────
# 1. 上传并解析简历
# ────────────────────────────────────────
@app.post("/api/upload")
async def upload_resume(file: UploadFile = File(...)):
    suffix = Path(file.filename).suffix.lower()
    if suffix not in (".pdf", ".doc", ".docx"):
        raise HTTPException(status_code=400, detail="仅支持 PDF 和 Word 格式（.pdf / .doc / .docx）")

    file_id = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{file_id}{suffix}"

    content = await file.read()
    save_path.write_bytes(content)

    resume_text = ""
    try:
        if suffix == ".pdf":
            with pdfplumber.open(str(save_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        resume_text += text + "\n"
        else:
            doc = DocxDocument(str(save_path))
            for para in doc.paragraphs:
                if para.text.strip():
                    resume_text += para.text + "\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败：{e}")

    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="无法从文件中提取文字，请确保文件内容不为空或非图片 PDF")

    return {
        "file_id": file_id,
        "filename": file.filename,
        "resume_text": resume_text,
        "char_count": len(resume_text),
    }


# ────────────────────────────────────────
# 2. 使用 GLM 分析简历
# ────────────────────────────────────────
@app.post("/api/analyze")
async def analyze_resume(
    resume_text: str = Form(...),
    jd_text: str    = Form(""),
    file_id: str    = Form(...),
):
    has_jd = bool(jd_text.strip())

    jd_block = f"""目标岗位职位描述（JD）：
{jd_text[:1500]}

请额外完成：
- matchScore：与该 JD 的匹配度评分（0-100）
- 结合 JD 给出更具针对性的 issues 建议
""" if has_jd else "（用户未提供目标 JD，请仅按大厂产品经理通用标准评估）"

    prompt = f"""你是一位有 10 年以上互联网大厂（阿里、腾讯、字节、美团）招聘经验的产品经理简历评估专家。

请对以下简历进行专业评估，并严格以合法 JSON 格式返回结果，不要输出任何 markdown 代码块标记。

【待评估简历】
{resume_text[:3500]}

{jd_block}

【评估维度（大厂 PM 标准）】
1. 结构清晰度（排版逻辑）
2. STAR 法则应用（情境-任务-行动-结果）
3. 数据量化程度（核心成果有无具体数字）
4. 产品经理关键词密度
5. 项目亮点与核心竞争力
6. 职业目标清晰度

【返回 JSON 格式】
{{
  "resumeScore": 数字(0-100),
  "matchScore": 数字(0-100)或null,
  "summary": "整体评价，2~3句话",
  "scoreDetails": {{
    "structure": 数字(0-20),
    "starMethod": 数字(0-20),
    "dataQuantification": 数字(0-20),
    "keywords": 数字(0-20),
    "highlights": 数字(0-20)
  }},
  "issues": [
    {{
      "id": "1",
      "category": "问题分类（如：数据量化不足）",
      "title": "问题标题（简短，10字以内）",
      "description": "问题详细说明",
      "suggestion": "具体优化建议（含示例改写）",
      "priority": "high",
      "section": "涉及板块（如：工作经历）"
    }}
  ]
}}

注意：
- issues 数组包含 4~6 个最重要的优化点
- priority 枚举：high / medium / low
- 所有文字使用中文
- 只返回 JSON，不要任何其他文字"""

    try:
        client = ZhipuAI(api_key=API_KEY)
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=3000,
        )
        raw = resp.choices[0].message.content.strip()

        # 清理 markdown 代码围栏
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]
        raw = raw.strip()

        result = json.loads(raw)
        result["file_id"] = file_id
        result["has_jd"]  = has_jd
        return result

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"AI 返回格式异常，请重试：{e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"分析失败：{e}")


# ────────────────────────────────────────
# 3. 生成优化简历
# ────────────────────────────────────────
@app.post("/api/generate")
async def generate_optimized(
    resume_text:      str = Form(...),
    jd_text:          str = Form(""),
    selected_issues:  str = Form("[]"),
    file_id:          str = Form(...),
):
    selected = json.loads(selected_issues)
    issues_desc = "\n".join(f"- {s}" for s in selected) if selected else "按大厂 PM 简历最佳实践全面优化"

    jd_block = f"【目标岗位 JD】\n{jd_text[:1000]}\n" if jd_text.strip() else ""

    prompt = f"""你是顶级产品经理简历写作专家，精通大厂 PM 简历包装技巧。

【原始简历】
{resume_text[:3500]}

{jd_block}
【需要采纳的优化点】
{issues_desc}

请生成优化后的完整简历，严格以合法 JSON 格式返回，不含任何 markdown 标记。

【返回 JSON 格式】
{{
  "name": "姓名",
  "title": "职位头衔（如：高级产品经理）",
  "contact": {{
    "email": "邮箱",
    "phone": "电话",
    "location": "城市",
    "wechat": "微信（可选）"
  }},
  "summary": "个人亮点/职业概述（3句话，突出数字成果与核心竞争力）",
  "workExperience": [
    {{
      "company": "公司名称",
      "department": "部门",
      "title": "职位",
      "period": "2021.03 - 至今",
      "location": "城市",
      "achievements": [
        "用 STAR 法则改写的成就句（含具体数据）",
        "成就 2"
      ]
    }}
  ],
  "education": [
    {{
      "school": "学校",
      "major": "专业",
      "degree": "学位",
      "period": "2015.09 - 2019.07"
    }}
  ],
  "projects": [
    {{
      "name": "项目名称",
      "role": "担任角色",
      "period": "时间",
      "description": "项目背景简介",
      "achievements": ["成果1（含数据）", "成果2"]
    }}
  ],
  "skills": {{
    "数据分析": ["SQL", "Python", "Tableau"],
    "产品工具": ["Axure", "Figma", "PRD"],
    "方法论": ["敏捷开发", "用户研究", "A/B 测试"]
  }}
}}

优化要求：
1. 所有工作/项目成就必须加入具体数据指标（百分比、倍数、金额等）
2. 使用 STAR 法则重写工作经历条目
3. 突出 PM 核心能力（数据驱动、需求管理、跨团队协作）
4. 语言风格简洁专业，符合大厂 JD 语境
5. 只做表述优化，不编造不存在的经历
6. 全部使用中文"""

    try:
        client = ZhipuAI(api_key=API_KEY)
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=4000,
        )
        raw = resp.choices[0].message.content.strip()

        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]
        raw = raw.strip()

        resume_data = json.loads(raw)

        # 生成 PDF
        pdf_path  = OUTPUT_DIR / f"{file_id}_optimized.pdf"
        word_path = OUTPUT_DIR / f"{file_id}_optimized.docx"
        _generate_pdf(resume_data, str(pdf_path))
        _generate_word(resume_data, str(word_path))

        return {
            "success":     True,
            "resume_data": resume_data,
            "pdf_url":     f"/api/download/{file_id}/pdf",
            "word_url":    f"/api/download/{file_id}/word",
        }

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"AI 返回格式异常：{e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成失败：{e}")


# ────────────────────────────────────────
# 4. 文件下载
# ────────────────────────────────────────
@app.get("/api/download/{file_id}/{fmt}")
async def download_file(file_id: str, fmt: str):
    if fmt == "pdf":
        path       = OUTPUT_DIR / f"{file_id}_optimized.pdf"
        media_type = "application/pdf"
        fname_en   = "resume_optimized.pdf"
        fname_zh   = "优化简历.pdf"
    elif fmt == "word":
        path       = OUTPUT_DIR / f"{file_id}_optimized.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        fname_en   = "resume_optimized.docx"
        fname_zh   = "优化简历.docx"
    else:
        raise HTTPException(status_code=400, detail="不支持的格式")

    if not path.exists():
        raise HTTPException(status_code=404, detail="文件尚未生成，请先执行生成操作")

    # RFC 5987 编码，浏览器显示中文文件名
    from urllib.parse import quote
    from starlette.responses import Response
    encoded = quote(fname_zh, safe="")
    content_disposition = (
        f'attachment; filename="{fname_en}"; '
        f"filename*=UTF-8''{encoded}"
    )
    data = path.read_bytes()
    return Response(
        content=data,
        media_type=media_type,
        headers={"Content-Disposition": content_disposition},
    )




# ────────────────────────────────────────
# PDF 生成工具函数
# ────────────────────────────────────────
def _generate_pdf(data: dict, output_path: str):
    COLOR_DARK = colors.HexColor("#1a1a2e")
    COLOR_BLUE = colors.HexColor("#2563eb")
    COLOR_GRAY = colors.HexColor("#64748b")
    COLOR_LB   = colors.HexColor("#eff6ff")  # light blue

    def _style(name, **kwargs):
        base = dict(fontName=CH_FONT, leading=16)
        base.update(kwargs)
        return ParagraphStyle(name, **base)

    S_NAME    = _style("name",    fontSize=22, leading=28, textColor=COLOR_DARK, spaceAfter=2)
    S_TITLE   = _style("title",   fontSize=12, textColor=COLOR_BLUE, spaceAfter=2)
    S_CONTACT = _style("contact", fontSize=9,  textColor=COLOR_GRAY, spaceAfter=6)
    S_SECTION = _style("section", fontSize=11, textColor=COLOR_BLUE, spaceBefore=14, spaceAfter=4, leading=16)
    S_BODY    = _style("body",    fontSize=9,  textColor=COLOR_DARK, leading=15)
    S_BULLET  = _style("bullet",  fontSize=9,  textColor=COLOR_DARK, leading=15, leftIndent=12)
    S_COMPANY = _style("company", fontSize=10, textColor=COLOR_DARK, leading=15)
    S_META    = _style("meta",    fontSize=9,  textColor=COLOR_GRAY, leading=14)

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=1.8*cm, bottomMargin=1.8*cm,
    )

    story = []

    # ── 头部
    story.append(Paragraph(data.get("name", ""), S_NAME))
    story.append(Paragraph(data.get("title", ""), S_TITLE))

    contact = data.get("contact", {})
    cparts = []
    if contact.get("email"):    cparts.append(f"✉ {contact['email']}")
    if contact.get("phone"):    cparts.append(f"✆ {contact['phone']}")
    if contact.get("location"): cparts.append(f"⊙ {contact['location']}")
    if contact.get("wechat"):   cparts.append(f"WeChat: {contact['wechat']}")
    story.append(Paragraph("   |   ".join(cparts), S_CONTACT))
    story.append(HRFlowable(width="100%", thickness=1.5, color=COLOR_BLUE, spaceBefore=2, spaceAfter=10))

    # ── 职业亮点
    summary = data.get("summary", "")
    if summary:
        story.append(Paragraph("职业亮点", S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_LB, spaceAfter=4))
        story.append(Paragraph(summary, S_BODY))

    # ── 工作经历
    work_exp = data.get("workExperience", [])
    if work_exp:
        story.append(Paragraph("工作经历", S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_LB, spaceAfter=4))
        for exp in work_exp:
            company = exp.get("company", "")
            dept    = exp.get("department", "")
            ttl     = exp.get("title", "")
            period  = exp.get("period", "")
            loc     = exp.get("location", "")

            comp_line = f"<b>{company}</b>  {dept}"
            meta_line = f"{ttl}   {period}   {loc}"

            block = [
                Paragraph(comp_line, S_COMPANY),
                Paragraph(meta_line, S_META),
            ]
            for ach in exp.get("achievements", []):
                block.append(Paragraph(f"• {ach}", S_BULLET))
            block.append(Spacer(1, 6))
            story.append(KeepTogether(block))

    # ── 项目经历
    projects = data.get("projects", [])
    if projects:
        story.append(Paragraph("项目经历", S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_LB, spaceAfter=4))
        for proj in projects:
            pname  = proj.get("name", "")
            prole  = proj.get("role", "")
            ptime  = proj.get("period", "")
            pdesc  = proj.get("description", "")

            block = [
                Paragraph(f"<b>{pname}</b>   {prole}   {ptime}", S_COMPANY),
            ]
            if pdesc:
                block.append(Paragraph(pdesc, S_META))
            for ach in proj.get("achievements", []):
                block.append(Paragraph(f"• {ach}", S_BULLET))
            block.append(Spacer(1, 6))
            story.append(KeepTogether(block))

    # ── 教育背景
    edu_list = data.get("education", [])
    if edu_list:
        story.append(Paragraph("教育背景", S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_LB, spaceAfter=4))
        for edu in edu_list:
            school = edu.get("school", "")
            major  = edu.get("major", "")
            degree = edu.get("degree", "")
            period = edu.get("period", "")
            story.append(Paragraph(f"<b>{school}</b>   {major}  {degree}   {period}", S_BODY))
            story.append(Spacer(1, 4))

    # ── 核心技能
    skills = data.get("skills", {})
    if skills:
        story.append(Paragraph("核心技能", S_SECTION))
        story.append(HRFlowable(width="100%", thickness=0.5, color=COLOR_LB, spaceAfter=4))
        table_data = []
        cats = list(skills.items())
        row = []
        for i, (cat, items) in enumerate(cats):
            cell_text = f"<b>{cat}</b>\n" + "  /  ".join(items)
            row.append(Paragraph(cell_text, S_BODY))
            if len(row) == 3 or i == len(cats) - 1:
                while len(row) < 3:
                    row.append(Paragraph("", S_BODY))
                table_data.append(row)
                row = []

        if table_data:
            col_w = (A4[0] - 4*cm) / 3
            tbl = Table(table_data, colWidths=[col_w]*3)
            tbl.setStyle(TableStyle([
                ("VALIGN",    (0,0), (-1,-1), "TOP"),
                ("TOPPADDING",(0,0), (-1,-1), 4),
                ("BOTTOMPADDING",(0,0), (-1,-1), 4),
            ]))
            story.append(tbl)

    doc.build(story)


# ────────────────────────────────────────
# Word 生成工具函数
# ────────────────────────────────────────
def _generate_word(data: dict, output_path: str):
    doc = WordDocument()

    # 页边距
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    BLUE  = RGBColor(0x25, 0x63, 0xEB)
    DARK  = RGBColor(0x1a, 0x1a, 0x2e)
    GRAY  = RGBColor(0x64, 0x74, 0x8b)

    def add_para(text, bold=False, size=10, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        return p

    def add_hr():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2563EB")
        pb.append(bot)
        pr.append(pb)

    def add_section_title(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size  = Pt(12)
        run.font.color.rgb = BLUE

    # 头部
    add_para(data.get("name", ""),  bold=True, size=20, color=DARK, space_after=2)
    add_para(data.get("title", ""), bold=False, size=11, color=BLUE, space_after=4)

    contact = data.get("contact", {})
    cparts = []
    if contact.get("email"):    cparts.append(f"✉ {contact['email']}")
    if contact.get("phone"):    cparts.append(f"✆ {contact['phone']}")
    if contact.get("location"): cparts.append(f"⊙ {contact['location']}")
    add_para("   |   ".join(cparts), size=9, color=GRAY, space_after=6)
    add_hr()

    # 职业亮点
    summary = data.get("summary", "")
    if summary:
        add_section_title("职业亮点")
        add_hr()
        add_para(summary, size=10, color=DARK)

    # 工作经历
    work_exp = data.get("workExperience", [])
    if work_exp:
        add_section_title("工作经历")
        add_hr()
        for exp in work_exp:
            add_para(
                f"{exp.get('company','')}  {exp.get('department','')}",
                bold=True, size=10, color=DARK, space_after=1
            )
            add_para(
                f"{exp.get('title','')}  {exp.get('period','')}  {exp.get('location','')}",
                size=9, color=GRAY, space_after=2
            )
            for ach in exp.get("achievements", []):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(ach)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

    # 项目经历
    projects = data.get("projects", [])
    if projects:
        add_section_title("项目经历")
        add_hr()
        for proj in projects:
            add_para(
                f"{proj.get('name','')}  {proj.get('role','')}  {proj.get('period','')}",
                bold=True, size=10, color=DARK, space_after=1
            )
            if proj.get("description"):
                add_para(proj["description"], size=9, color=GRAY, space_after=2)
            for ach in proj.get("achievements", []):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(ach)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

    # 教育背景
    edu_list = data.get("education", [])
    if edu_list:
        add_section_title("教育背景")
        add_hr()
        for edu in edu_list:
            add_para(
                f"{edu.get('school','')}  {edu.get('major','')}  {edu.get('degree','')}  {edu.get('period','')}",
                size=10, color=DARK
            )

    # 核心技能
    skills = data.get("skills", {})
    if skills:
        add_section_title("核心技能")
        add_hr()
        for cat, items in skills.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run_cat = p.add_run(f"{cat}：")
            run_cat.bold = True
            run_cat.font.size = Pt(9)
            run_cat.font.color.rgb = BLUE
            run_items = p.add_run("  /  ".join(items))
            run_items.font.size = Pt(9)
            run_items.font.color.rgb = DARK

    doc.save(output_path)


# ────────────────────────────────────────
# 启动
# ────────────────────────────────────────
if __name__ == "__main__":
    uvicorn.run("server:app", host="0.0.0.0", port=8001, reload=True)
